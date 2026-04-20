import os
import uuid
import re
from datetime import datetime

# ── Maximum characters extracted from an uploaded file ────────────────────────
_MAX_FILE_CHARS = 12_000


def extract_file_content(file_path: str) -> str:
    """
    Extracts plain text from an uploaded file so it can be injected into the
    LLM prompt for direct analysis.

    Supported formats: PDF, DOCX, XLSX, CSV, TXT, MD
    Returns an empty string on failure (with a note embedded).
    """
    ext = os.path.splitext(file_path)[1].lower()

    try:
        # ── PDF ──────────────────────────────────────────────────────────────
        if ext == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                try:
                    from PyPDF2 import PdfReader  # legacy fallback
                except ImportError:
                    return "[PDF extraction unavailable — pypdf not installed]"
            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return "\n\n".join(pages)[:_MAX_FILE_CHARS]

        # ── DOCX ─────────────────────────────────────────────────────────────
        elif ext == ".docx":
            try:
                from docx import Document
            except ImportError:
                return "[DOCX extraction unavailable — python-docx not installed]"
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)[:_MAX_FILE_CHARS]

        # ── XLSX ─────────────────────────────────────────────────────────────
        elif ext in (".xlsx", ".xls"):
            try:
                from openpyxl import load_workbook
            except ImportError:
                return "[XLSX extraction unavailable — openpyxl not installed]"
            wb = load_workbook(file_path, read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"=== Sheet: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [str(c) if c is not None else "" for c in row]
                    if any(v.strip() for v in row_vals):
                        lines.append("\t".join(row_vals))
            return "\n".join(lines)[:_MAX_FILE_CHARS]

        # ── CSV ──────────────────────────────────────────────────────────────
        elif ext == ".csv":
            import csv
            lines = []
            with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                for row in reader:
                    lines.append(", ".join(row))
            return "\n".join(lines)[:_MAX_FILE_CHARS]

        # ── TXT / MD ─────────────────────────────────────────────────────────
        elif ext in (".txt", ".md"):
            with open(file_path, encoding="utf-8", errors="replace") as f:
                return f.read()[:_MAX_FILE_CHARS]

        else:
            return f"[Unsupported file type: {ext}]"

    except Exception as exc:
        return f"[Error reading file: {exc}]"
    finally:
        # Final safety check: if the result looks like binary, clear it
        pass

def sanitize_text(text: str) -> str:
    """Removes non-printable characters and binary-like patterns."""
    if not text: return ""
    # Remove null bytes and other control chars
    text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
    # Remove common binary headers if they leaked
    if text.startswith("%PDF") or "stream" in text[:100]:
        return "[Error: Binary data detected and stripped]"
    return text[:_MAX_FILE_CHARS]


class WaslaToolKit:
    """
    Core toolset for document generation and structured output processing.
    Handles PDF, DOCX, Markdown and Plan generation.
    """
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def _safe_filename(self, filename: str, ext: str) -> str:
        # Remove potentially dangerous characters
        safe_name = re.sub(r'[^\w\s\.-]', '', filename).strip().replace(" ", "_")
        if not safe_name:
            safe_name = "generated_document"
        if not safe_name.lower().endswith(ext.lower()):
            safe_name += ext
        return f"{uuid.uuid4().hex[:8]}_{safe_name}"

    def generate_markdown(self, content: str, filename: str) -> str:
        """Saves content as a Markdown file."""
        unique_name = self._safe_filename(filename, ".md")
        path = os.path.join(self.output_dir, unique_name)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return unique_name

    def _setup_pdf(self, content: str) -> "FPDF":
        """Initializes FPDF with proper fonts and Unicode support if needed."""
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Check if content contains Arabic characters
        has_arabic = any('\u0600' <= c <= '\u06FF' for c in content)
        
        if has_arabic:
            # Try to load Amiri font for Arabic support
            font_paths = [
                "/usr/share/fonts/opentype/fonts-hosny-amiri/Amiri-Regular.ttf",
                "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf"
            ]
            loaded = False
            for path in font_paths:
                if os.path.exists(path):
                    try:
                        pdf.add_font("Amiri", "", path)
                        pdf.set_font("Amiri", size=12)
                        loaded = True
                        break
                    except: continue
            
            if not loaded:
                pdf.set_font("Arial", size=11)
        else:
            pdf.set_font("Arial", size=11)
            
        return pdf

    def generate_pdf(self, content: str, filename: str) -> str:
        """Converts Markdown/Text content to PDF using fpdf2."""
        try:
            from fpdf import FPDF
        except ImportError:
            return self.generate_markdown(f"PDF Generation failed (fpdf2 not installed). Content:\n\n{content}", filename + ".txt")

        unique_name = self._safe_filename(filename, ".pdf")
        path = os.path.join(self.output_dir, unique_name)
        
        pdf = self._setup_pdf(content)
        is_unicode = "Amiri" in pdf.fonts
        
        # Add a header
        if is_unicode:
            pdf.set_font("Amiri", size=18)
        else:
            pdf.set_font("Arial", 'B', 16)
            
        pdf.cell(0, 10, filename.replace(".pdf", "").replace(".docx", "").replace("_", " ").title(), ln=True, align='C')
        pdf.ln(5)
        
        # Process content
        if is_unicode:
            pdf.set_font("Amiri", size=12)
        else:
            pdf.set_font("Arial", size=11)
        
        # Simple cleanup
        clean_content = content.replace("**", "").replace("__", "").replace("#", "")
        
        for line in clean_content.split('\n'):
            if not is_unicode:
                line = line.encode('latin-1', 'replace').decode('latin-1')
            
            # Fix: Ensure x is at left margin and use epw
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(pdf.epw, 8, line)
            
        pdf.output(path)
        return unique_name

    def generate_docx(self, content: str, filename: str) -> str:
        """Converts content to a Word Document using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return self.generate_markdown(f"DOCX Generation failed (python-docx not installed). Content:\n\n{content}", filename + ".txt")

        unique_name = self._safe_filename(filename, ".docx")
        path = os.path.join(self.output_dir, unique_name)
        
        doc = Document()
        doc.add_heading(filename.replace(".docx", "").replace("_", " ").title(), 0)
        
        # Basic parsing for lines
        for line in content.split('\n'):
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line)
        
        doc.save(path)
        return unique_name

    def generate_plan(self, topic: str, details: str) -> str:
        """Specialized tool for creating a structured study plan."""
        # This is a helper that formats the input into a nice MD/PDF structure
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        plan_content = f"""# Study Plan: {topic}
Generated on: {timestamp}

## Overview
{details}

## Structured Timeline
| Phase | Focus Area | Estimated Time |
|-------|------------|----------------|
| 1     | Foundations| 2 Hours        |
| 2     | Deep Dive  | 4 Hours        |
| 3     | Practice   | 3 Hours        |
| 4     | Review     | 1 Hour         |

---
*Created by Wasla Master AI Assistant*
"""
        return self.generate_markdown(plan_content, f"Plan_{topic.replace(' ', '_')}")

    def generate_cv(self, content: str, filename: str) -> str:
        """Specialized tool for generating a professional CV/Resume."""
        try:
            from fpdf import FPDF
        except ImportError:
            return self.generate_markdown(f"CV Generation failed (fpdf2 not installed). Content:\n\n{content}", filename + ".txt")

        unique_name = self._safe_filename(filename, ".pdf")
        path = os.path.join(self.output_dir, unique_name)
        
        pdf = self._setup_pdf(content)
        is_unicode = "Amiri" in pdf.fonts
        
        # --- Professional Header ---
        pdf.set_fill_color(139, 0, 0) # Dark Red (Primary)
        pdf.rect(0, 0, 210, 40, 'F')
        
        pdf.set_text_color(255, 255, 255)
        if is_unicode:
            pdf.set_font("Amiri", size=24)
        else:
            pdf.set_font("Arial", 'B', 24)
        
        # Try to extract Name from first line
        lines = content.split('\n')
        name = "CURRICULUM VITAE"
        if lines and len(lines[0]) < 50:
            name = lines[0].strip().replace('#', '').upper()
            lines = lines[1:] # Remove name from body
            
        pdf.set_y(15)
        pdf.cell(0, 10, name, ln=True, align='C')
        
        if is_unicode:
            pdf.set_font("Amiri", size=10)
        else:
            pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, "Generated by Wasla AI Professional Career Suite", ln=True, align='C')
        
        # --- Body ---
        pdf.set_text_color(30, 30, 30)
        pdf.set_y(50)
        
        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(2)
                continue
                
            if line.startswith('### ') or line.startswith('## ') or line.startswith('# '):
                # Section Header
                section_name = line.replace('#', '').strip().upper()
                pdf.ln(5)
                if is_unicode:
                    pdf.set_font("Amiri", size=12)
                else:
                    pdf.set_font("Arial", 'B', 12)
                pdf.set_draw_color(139, 0, 0)
                pdf.set_line_width(0.5)
                pdf.cell(0, 8, section_name, ln=True)
                pdf.line(pdf.get_x(), pdf.get_y(), 200, pdf.get_y())
                pdf.ln(2)
            else:
                # Regular text
                if is_unicode:
                    pdf.set_font("Amiri", size=10)
                else:
                    pdf.set_font("Arial", size=10)
                # Handle basic bullet points
                if line.startswith('- ') or line.startswith('* '):
                    pdf.set_x(15)
                    line = "• " + line[2:]
                
                if not is_unicode:
                    line = line.encode('latin-1', 'replace').decode('latin-1')
                
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(pdf.epw, 6, line)
                
        pdf.output(path)
        return unique_name

    def generate_proposal(self, content: str, filename: str) -> str:
        """Specialized tool for generating a business project proposal."""
        try:
            from fpdf import FPDF
        except ImportError:
            return self.generate_markdown(f"Proposal Generation failed (fpdf2 not installed). Content:\n\n{content}", filename + ".txt")

        unique_name = self._safe_filename(filename, ".pdf")
        path = os.path.join(self.output_dir, unique_name)
        
        pdf = self._setup_pdf(content)
        is_unicode = "Amiri" in pdf.fonts
        
        # --- Professional Header ---
        # Dark Blue theme for Proposals
        pdf.set_fill_color(0, 51, 102) 
        pdf.rect(0, 0, 210, 45, 'F')
        
        pdf.set_text_color(255, 255, 255)
        if is_unicode:
            pdf.set_font("Amiri", size=22)
        else:
            pdf.set_font("Arial", 'B', 22)
        
        lines = content.split('\n')
        title = "PROJECT PROPOSAL"
        if lines and len(lines[0]) < 60:
            title = lines[0].strip().replace('#', '').upper()
            lines = lines[1:]
            
        pdf.set_y(15)
        pdf.cell(0, 10, title, ln=True, align='C')
        
        if is_unicode:
            pdf.set_font("Amiri", size=10)
        else:
            pdf.set_font("Arial", size=10)
        timestamp = datetime.now().strftime("%B %d, %Y")
        pdf.cell(0, 8, f"Date: {timestamp} | Generated by Wasla Master", ln=True, align='C')
        
        # --- Body ---
        pdf.set_text_color(30, 30, 30)
        pdf.set_y(55)
        
        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue
                
            if line.startswith('### ') or line.startswith('## ') or line.startswith('# '):
                section_name = line.replace('#', '').strip().upper()
                pdf.ln(6)
                if is_unicode:
                    pdf.set_font("Amiri", size=12)
                else:
                    pdf.set_font("Arial", 'B', 12)
                pdf.set_text_color(0, 51, 102)
                pdf.cell(0, 8, section_name, ln=True)
                # Subtle underline
                pdf.set_draw_color(200, 200, 200)
                pdf.set_line_width(0.3)
                pdf.line(pdf.get_x(), pdf.get_y(), 200, pdf.get_y())
                pdf.ln(3)
                pdf.set_text_color(30, 30, 30)
            else:
                if is_unicode:
                    pdf.set_font("Amiri", size=10)
                else:
                    pdf.set_font("Arial", size=10)
                if line.startswith('- ') or line.startswith('* '):
                    pdf.set_x(18)
                    line = "» " + line[2:] # different bullet style
                
                if not is_unicode:
                    line = line.encode('latin-1', 'replace').decode('latin-1')
                
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(pdf.epw, 6, line)
                
        pdf.output(path)
        return unique_name

    @staticmethod
    def extract_tags(text: str, tag_name: str) -> list:
        """
        Extracts content from specific XML-like tags.
        Supports attributes like filename="..." or name="..."
        """
        # Improved regex: matches <TAG attr="val">content</TAG> 
        # or <TAG attr="val">content (if closing tag is missing)
        # It looks for filename or name attributes specifically.
        pattern = rf'<{tag_name}(?:\s+(?:filename|name)="([^"]+)")?>([\s\S]*?)(?:</{tag_name}>|(?=<{tag_name}|\Z))'
        matches = []
        for match in re.finditer(pattern, text, re.IGNORECASE):
            content = match.group(2).strip()
            if not content: continue
            
            matches.append({
                "attr": match.group(1),
                "content": content,
                "raw": match.group(0)
            })
        return matches
